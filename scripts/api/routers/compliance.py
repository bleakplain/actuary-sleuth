"""合规检查路由 — 产品参数检查 + 条款文档审查 + 文档解析。"""

import os
import uuid
import json
import asyncio
import logging
import tempfile
from typing import Dict, List
from datetime import datetime

from fastapi import APIRouter, HTTPException, UploadFile, File, Form

from api.schemas.compliance import (
    ProductCheckRequest, DocumentCheckRequest, ComplianceReportOut,
    ParsedDocumentResponse, ParsedClause, ParsedPremiumTable, ParsedSection,
    RichTextParseRequest,
)
from api.dependencies import get_rag_engine

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/api/compliance", tags=["合规检查"])

_COMPLIANCE_PROMPT_PRODUCT = """你是一位保险法规合规专家。请根据以下产品参数和相关法规条款，逐项检查该产品是否符合法规要求。

## 产品信息
- 产品名称：{product_name}
- 险种类型：{category}
- 产品参数：{params_json}

## 相关法规条款
{context}

## 输出要求
请以 JSON 格式输出检查结果，严格遵循以下结构：
{{
    "summary": {{
        "compliant": <合规项数>,
        "non_compliant": <不合规项数>,
        "attention": <需关注项数>
    }},
    "items": [
        {{
            "param": "<参数名称>",
            "value": "<产品实际值>",
            "requirement": "<法规要求，引用法规原文关键句>",
            "status": "<compliant|non_compliant|attention>",
            "source": "<法规来源，格式：[来源X]>",
            "source_excerpt": "<从来源法规中直接摘录的原文片段，作为该判断的事实依据>",
            "suggestion": "<修改建议，仅不合规时填写>"
        }}
    ]
}}

注意：
1. 每个参数都要检查，未找到明确法规要求的标注为 attention
2. source 必须使用 [来源X] 格式引用法规条款
3. source_excerpt 必须是从对应来源中直接摘录的原文，不得自行编造或改写
4. requirement 应结合法规原文表述，使合规判断有据可查
5. 仅输出 JSON，不要附加其他文字
"""

_COMPLIANCE_PROMPT_DOCUMENT = """你是一位保险法规合规专家。请审查以下保险条款文档，检查是否符合相关法规要求。

## 条款文档内容
{document_content}

## 相关法规条款
{context}

## 输出要求
请以 JSON 格式输出检查结果，严格遵循以下结构：
{{
    "summary": {{
        "compliant": <合规项数>,
        "non_compliant": <不合规项数>,
        "attention": <需关注项数>
    }},
    "items": [
        {{
            "param": "<检查项名称>",
            "value": "<条款中的实际内容>",
            "requirement": "<法规要求，引用法规原文关键句>",
            "status": "<compliant|non_compliant|attention>",
            "source": "<法规来源，格式：[来源X]>",
            "source_excerpt": "<从来源法规中直接摘录的原文片段，作为该判断的事实依据>",
            "suggestion": "<修改建议>"
        }}
    ],
    "extracted_params": {{
        "<参数名>": "<提取值>"
    }}
}}

注意：
1. 先提取条款中的关键参数，再逐项检查合规性
2. 检查项包括但不限于：等待期、免赔额、保险期间、缴费方式、免责条款等
3. source 必须使用 [来源X] 格式引用法规条款
4. source_excerpt 必须是从对应来源中直接摘录的原文，不得自行编造或改写
5. requirement 应结合法规原文表述，使合规判断有据可查
6. 仅输出 JSON，不要附加其他文字
"""



def _build_context(search_results: list) -> str:
    parts = []
    for i, r in enumerate(search_results):
        law_name = r.get('law_name', '')
        article = r.get('article_number', '')
        content = r.get('content', '')
        authority = r.get('issuing_authority', '')
        doc_number = r.get('doc_number', '')
        effective = r.get('effective_date', '')
        header = f"[来源{i+1}] 【{law_name}】{article}"
        if doc_number:
            header += f"（{doc_number}）"
        if authority:
            header += f"\n发布机关：{authority}"
        if effective:
            header += f"\n生效日期：{effective}"
        parts.append(f"{header}\n{content}")
    return "\n\n".join(parts)


def _run_compliance_check(engine, prompt: str) -> Dict:
    result = engine.ask(prompt)
    answer = result.get("answer", "")

    try:
        json_start = answer.find("{")
        json_end = answer.rfind("}") + 1
        if json_start >= 0 and json_end > json_start:
            parsed = json.loads(answer[json_start:json_end])
        else:
            parsed = {"summary": {"compliant": 0, "non_compliant": 0, "attention": 0}, "items": []}
    except json.JSONDecodeError:
        parsed = {
            "summary": {"compliant": 0, "non_compliant": 0, "attention": 0},
            "items": [],
            "raw_answer": answer,
        }

    parsed["sources"] = result.get("sources", [])
    parsed["citations"] = result.get("citations", [])
    return parsed


@router.post("/check/product", response_model=ComplianceReportOut)
async def check_product(req: ProductCheckRequest):
    engine = get_rag_engine()

    query = f"{req.category} 保险产品合规要求 {req.product_name}"
    search_results = engine.search(query, top_k=10)

    context = _build_context(search_results)

    prompt = _COMPLIANCE_PROMPT_PRODUCT.format(
        product_name=req.product_name,
        category=req.category,
        params_json=json.dumps(req.params, ensure_ascii=False),
        context=context,
    )

    try:
        result = await asyncio.to_thread(_run_compliance_check, engine, prompt)
    except Exception as e:
        logger.error(f"Compliance check failed: {e}")
        raise HTTPException(status_code=500, detail=f"合规检查失败: {e}")

    report_id = f"cr_{uuid.uuid4().hex[:8]}"
    from api.database import save_compliance_report
    save_compliance_report(report_id, req.product_name, req.category, "product", result)

    return ComplianceReportOut(
        id=report_id,
        product_name=req.product_name,
        category=req.category,
        mode="product",
        result=result,
        created_at="",
    )


@router.post("/check/document", response_model=ComplianceReportOut)
async def check_document(req: DocumentCheckRequest):
    engine = get_rag_engine()

    try:
        from lib.llm.factory import LLMClientFactory
        llm = LLMClientFactory.create_qa_llm()
        extract_prompt = f"请从以下保险条款中提取关键参数（险种类型、等待期、免赔额等），以 JSON 格式输出：\n\n{req.document_content[:3000]}"
        extracted = llm.chat([{"role": "user", "content": extract_prompt}])
    except Exception:
        extracted = ""

    query = f"保险合规要求 {extracted[:200]}"
    search_results = engine.search(query, top_k=10)

    context = _build_context(search_results)

    prompt = _COMPLIANCE_PROMPT_DOCUMENT.format(
        document_content=req.document_content[:5000],
        context=context,
    )

    try:
        result = await asyncio.to_thread(_run_compliance_check, engine, prompt)
    except Exception as e:
        logger.error(f"Document check failed: {e}")
        raise HTTPException(status_code=500, detail=f"条款审查失败: {e}")

    report_id = f"cr_{uuid.uuid4().hex[:8]}"
    product_name = req.product_name or "未命名产品"
    from api.database import save_compliance_report
    save_compliance_report(report_id, product_name, "", "document", result)

    return ComplianceReportOut(
        id=report_id,
        product_name=product_name,
        category="",
        mode="document",
        result=result,
        created_at="",
    )


@router.get("/reports", response_model=list[ComplianceReportOut])
async def list_compliance_reports():
    from api.database import get_compliance_reports
    return get_compliance_reports()


@router.get("/reports/{report_id}", response_model=ComplianceReportOut)
async def get_compliance_report(report_id: str):
    from api.database import get_compliance_report
    report = get_compliance_report(report_id)
    if report is None:
        raise HTTPException(status_code=404, detail="报告不存在")
    return report


@router.delete("/reports/{report_id}")
async def delete_compliance_report(report_id: str):
    from api.database import get_connection
    with get_connection() as conn:
        cur = conn.execute(
            "DELETE FROM compliance_reports WHERE id = ?", (report_id,)
        )
        if cur.rowcount == 0:
            raise HTTPException(status_code=404, detail="报告不存在")
    return {"status": "deleted"}


ALLOWED_EXTENSIONS = ['.pdf', '.docx']


def _build_combined_text(
    clauses: List[ParsedClause],
    premium_tables: List[ParsedPremiumTable],
    notices: List[ParsedSection],
    health_disclosures: List[ParsedSection],
    exclusions: List[ParsedSection],
    rider_clauses: List[ParsedClause],
) -> str:
    """将解析结果合并为文本，用于合规检查"""
    parts = []
    for c in clauses:
        parts.append(f"【条款 {c.number}】{c.title}\n{c.text}")
    for i, t in enumerate(premium_tables, 1):
        parts.append(f"【费率表 {i}】\n{t.raw_text}")
    for s in notices:
        parts.append(f"【投保须知】{s.title}\n{s.content}")
    for s in health_disclosures:
        parts.append(f"【健康告知】{s.title}\n{s.content}")
    for s in exclusions:
        parts.append(f"【责任免除】{s.title}\n{s.content}")
    for c in rider_clauses:
        parts.append(f"【附加险条款 {c.number}】{c.title}\n{c.text}")
    return "\n\n".join(parts)


def _audit_doc_to_response(audit_doc, file_type: str) -> ParsedDocumentResponse:
    """将 AuditDocument 转换为 ParsedDocumentResponse"""
    clauses = [
        ParsedClause(number=c.number, title=c.title, text=c.text)
        for c in audit_doc.clauses
    ]
    tables = [
        ParsedPremiumTable(
            table_type=t.table_type.value if hasattr(t.table_type, 'value') else str(t.table_type),
            remark=t.remark or "",
            raw_text=t.raw_text,
            data=[list(row) for row in t.data]
        )
        for t in audit_doc.tables
    ]
    notices = [
        ParsedSection(title=s.title, content=s.content)
        for s in audit_doc.notices
    ]
    health_disclosures = [
        ParsedSection(title=s.title, content=s.content)
        for s in audit_doc.health_disclosures
    ]
    exclusions = [
        ParsedSection(title=s.title, content=s.content)
        for s in audit_doc.exclusions
    ]
    rider_clauses = [
        ParsedClause(number=c.number, title=c.title, text=c.text)
        for c in audit_doc.rider_clauses
    ]

    combined_text = _build_combined_text(
        clauses, tables, notices, health_disclosures, exclusions, rider_clauses
    )

    return ParsedDocumentResponse(
        parse_id=f"pd_{uuid.uuid4().hex[:8]}",
        file_name=audit_doc.file_name,
        file_type=file_type,
        clauses=clauses,
        premium_tables=tables,
        notices=notices,
        health_disclosures=health_disclosures,
        exclusions=exclusions,
        rider_clauses=rider_clauses,
        warnings=list(audit_doc.warnings),
        combined_text=combined_text,
        parse_time=audit_doc.parse_time.isoformat(),
    )


@router.post("/parse-file", response_model=ParsedDocumentResponse)
async def parse_file(file: UploadFile = File(...)):
    """上传并解析保险产品文档（PDF/DOCX）"""
    ext = os.path.splitext(file.filename or "")[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail=f"仅支持 PDF 和 DOCX 格式，当前: {ext}")

    with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
        content = await file.read()
        tmp.write(content)
        tmp_path = tmp.name

    try:
        from lib.doc_parser import parse_product_document, DocumentParseError
        audit_doc = parse_product_document(tmp_path)
        return _audit_doc_to_response(audit_doc, ext)
    except DocumentParseError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"Document parse failed: {e}")
        raise HTTPException(status_code=500, detail=f"文档解析失败: {e}")
    finally:
        os.unlink(tmp_path)


def _html_to_docx(html_content: str) -> str:
    """将 HTML 转换为临时 DOCX 文件，返回文件路径"""
    from docx import Document
    from html.parser import HTMLParser
    import re

    class SimpleHTMLParser(HTMLParser):
        def __init__(self):
            super().__init__()
            self.paragraphs = []
            self.tables = []
            self.current_table = []
            self.current_row = []
            self.current_cell = ""
            self.in_table = False
            self.in_row = False
            self.in_cell = False
            self.current_text = ""
            self.in_p = False

        def handle_starttag(self, tag, attrs):
            if tag == 'p':
                self.in_p = True
                self.current_text = ""
            elif tag == 'table':
                self.in_table = True
                self.current_table = []
            elif tag == 'tr':
                self.in_row = True
                self.current_row = []
            elif tag in ['td', 'th']:
                self.in_cell = True
                self.current_cell = ""
            elif tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                self.in_p = True
                self.current_text = ""

        def handle_endtag(self, tag):
            if tag == 'p' and self.in_p:
                self.in_p = False
                text = self.current_text.strip()
                if text:
                    self.paragraphs.append(text)
            elif tag == 'table':
                self.in_table = False
                if self.current_table:
                    self.tables.append(self.current_table)
                self.current_table = []
            elif tag == 'tr':
                self.in_row = False
                if self.current_row:
                    self.current_table.append(self.current_row)
                self.current_row = []
            elif tag in ['td', 'th']:
                self.in_cell = False
                self.current_row.append(self.current_cell.strip())
            elif tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                self.in_p = False
                text = self.current_text.strip()
                if text:
                    self.paragraphs.append(text)

        def handle_data(self, data):
            if self.in_p:
                self.current_text += data
            elif self.in_cell:
                self.current_cell += data

    parser = SimpleHTMLParser()
    parser.feed(html_content)

    doc = Document()

    for p in parser.paragraphs:
        doc.add_paragraph(p)

    for table_data in parser.tables:
        if not table_data or not table_data[0]:
            continue
        num_cols = max(len(row) for row in table_data)
        table = doc.add_table(rows=len(table_data), cols=num_cols)
        table.style = 'Table Grid'
        for i, row in enumerate(table_data):
            for j, cell in enumerate(row):
                if j < num_cols:
                    table.rows[i].cells[j].text = cell

    tmp_path = tempfile.mktemp(suffix='.docx')
    doc.save(tmp_path)
    return tmp_path


@router.post("/parse-rich-text", response_model=ParsedDocumentResponse)
async def parse_rich_text(req: RichTextParseRequest):
    """解析富文本内容（HTML）"""
    from lib.doc_parser import DocumentParseError

    if not req.html_content or not req.html_content.strip():
        raise HTTPException(status_code=400, detail="HTML 内容不能为空")

    tmp_path = None
    try:
        tmp_path = _html_to_docx(req.html_content)
        from lib.doc_parser import parse_product_document
        audit_doc = parse_product_document(tmp_path)
        response = _audit_doc_to_response(audit_doc, ".html")
        if req.product_name:
            response.file_name = req.product_name
        return response
    except DocumentParseError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"Rich text parse failed: {e}")
        raise HTTPException(status_code=500, detail=f"富文本解析失败: {e}")
    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)
