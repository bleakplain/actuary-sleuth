#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Word 文档解析器"""
from __future__ import annotations

import logging
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any, Optional

from docx import Document
from docx.table import Table

import re

from ..models import Clause, DataTable, AuditDocument, DocumentParseError, SectionType, TableType
from .section_detector import SectionDetector
from .table_classifier import TableClassifier
from .utils import split_title_and_content, add_section

logger = logging.getLogger(__name__)

# 条款编号格式：只匹配 X.Y 或 X.Y.Z 格式（至少包含一个点）
CLAUSE_NUMBER_PATTERN = re.compile(r'^\d+\.\d+(?:\.\d+)*$')


class DocxParser:
    """Word (.docx) 文档解析器"""

    def __init__(self, section_detector: Optional[SectionDetector] = None):
        self.detector = section_detector or SectionDetector()
        self.table_classifier = TableClassifier()

    @staticmethod
    def supported_extensions() -> List[str]:
        return ['.docx']

    def parse(self, file_path: str) -> AuditDocument:
        path = Path(file_path)
        if not path.exists():
            raise DocumentParseError("文件不存在", file_path)

        try:
            doc = Document(file_path)
        except Exception as e:
            raise DocumentParseError("Word 文件解析失败", file_path, str(e))

        warnings: List[str] = []

        clauses = self._extract_clauses(doc.tables, warnings)
        tables = self._extract_tables(doc.tables, warnings)
        sections = self._extract_sections(doc.paragraphs, warnings)

        return AuditDocument(
            file_name=path.name,
            file_type='.docx',
            clauses=clauses,
            tables=tables,
            notices=sections['notices'],
            health_disclosures=sections['health_disclosures'],
            exclusions=sections['exclusions'],
            rider_clauses=sections['rider_clauses'],
            parse_time=datetime.now(),
            warnings=warnings,
        )

    def _extract_clauses(self, tables: List[Table], warnings: List[str]) -> List[Clause]:
        """提取条款，只提取 X.Y 格式的条款，过滤章节标题。

        支持多种表格结构：
        - 2 列：编号 + 内容（标题和正文合并）
        - 3+ 列：编号 + 标题 + 正文（正文在第 3 列）
        """
        clauses = []

        for table in tables:
            if not table.rows:
                continue

            first_row = [cell.text.strip() for cell in table.rows[0].cells]
            if self.detector.is_non_clause_table(first_row):
                continue

            # 判断表格列结构
            num_cols = len(table.rows[0].cells) if table.rows else 0

            for row in table.rows:
                cells = [(cell.text or '').strip() for cell in row.cells]
                if not cells or not cells[0]:
                    continue

                number = cells[0].strip()
                # 只匹配 X.Y 或 X.Y.Z 格式，过滤单数字章节标题
                if CLAUSE_NUMBER_PATTERN.match(number):
                    if num_cols >= 3:
                        # 3+ 列结构：编号 + 标题 + 正文
                        title = cells[1] if len(cells) > 1 else ''
                        text = cells[2] if len(cells) > 2 else ''
                    else:
                        # 2 列结构：编号 + 内容（需分离标题和正文）
                        title, text = split_title_and_content(cells[1] if len(cells) > 1 else '')
                    clauses.append(Clause(number=number, title=title, text=text))

        return clauses

    def _extract_tables(self, tables: List[Table], warnings: List[str]) -> List[DataTable]:
        """提取数据表格，使用 TableClassifier 分类表格类型。"""
        result: List[DataTable] = []

        for table in tables:
            if not table.rows:
                continue

            rows = [[(cell.text or '').strip() for cell in row.cells] for row in table.rows]
            if len(rows) < 2:
                continue

            header = [h for h in rows[0] if h]
            if len(header) < 2:
                continue

            classification = self.table_classifier.classify_by_header(' '.join(rows[0]))
            raw_text = '\n'.join('\t'.join(row) for row in rows)

            result.append(DataTable(
                data=rows,
                table_type=classification,
                raw_text=raw_text,
            ))

        return result

    def _extract_sections(
        self,
        paragraphs: List,
        warnings: List[str],
    ) -> Dict[str, List[Any]]:
        result: Dict[str, List[Any]] = {
            'notices': [],
            'health_disclosures': [],
            'exclusions': [],
            'rider_clauses': [],
        }

        current_type: Optional[SectionType] = None
        current_content: List[str] = []

        for para in paragraphs:
            text = para.text.strip()
            if not text:
                continue

            detected = self.detector.detect_section_type(text)
            if detected:
                if current_type and current_content:
                    add_section(result, current_type, '', '\n'.join(current_content))
                current_type = detected
                current_content = []
            else:
                if current_type:
                    current_content.append(text)

        if current_type and current_content:
            add_section(result, current_type, '', '\n'.join(current_content))

        return result
