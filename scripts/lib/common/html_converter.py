"""HTML 转 DOCX 转换工具"""
import os
import tempfile
from html.parser import HTMLParser


class SimpleHTMLParser(HTMLParser):
    def __init__(self):
        super().__init__()
        self.paragraphs = []
        self.tables = []
        self.current_table = []
        self.current_row = []
        self.current_cell = ""
        self.in_table = False
        self.in_row = False
        self.in_cell = False
        self.current_text = ""
        self.in_p = False

    def handle_starttag(self, tag, attrs):
        if tag == 'p':
            self.in_p = True
            self.current_text = ""
        elif tag == 'table':
            self.in_table = True
            self.current_table = []
        elif tag == 'tr':
            self.in_row = True
            self.current_row = []
        elif tag in ['td', 'th']:
            self.in_cell = True
            self.current_cell = ""
        elif tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            self.in_p = True
            self.current_text = ""

    def handle_endtag(self, tag):
        if tag == 'p' and self.in_p:
            self.in_p = False
            text = self.current_text.strip()
            if text:
                self.paragraphs.append(text)
        elif tag == 'table':
            self.in_table = False
            if self.current_table:
                self.tables.append(self.current_table)
            self.current_table = []
        elif tag == 'tr':
            self.in_row = False
            if self.current_row:
                self.current_table.append(self.current_row)
            self.current_row = []
        elif tag in ['td', 'th']:
            self.in_cell = False
            self.current_row.append(self.current_cell.strip())
        elif tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            self.in_p = False
            text = self.current_text.strip()
            if text:
                self.paragraphs.append(text)

    def handle_data(self, data):
        if self.in_p:
            self.current_text += data
        elif self.in_cell:
            self.current_cell += data


def html_to_docx(html_content: str) -> str:
    """将 HTML 转换为临时 DOCX 文件，返回文件路径"""
    from docx import Document

    parser = SimpleHTMLParser()
    parser.feed(html_content)

    doc = Document()
    for p in parser.paragraphs:
        doc.add_paragraph(p)
    for table_data in parser.tables:
        if not table_data or not table_data[0]:
            continue
        num_cols = max(len(row) for row in table_data)
        table = doc.add_table(rows=len(table_data), cols=num_cols)
        table.style = 'Table Grid'
        for i, row in enumerate(table_data):
            for j, cell in enumerate(row):
                if j < num_cols:
                    table.rows[i].cells[j].text = cell

    fd, tmp_path = tempfile.mkstemp(suffix='.docx')
    os.close(fd)
    doc.save(tmp_path)
    return tmp_path
