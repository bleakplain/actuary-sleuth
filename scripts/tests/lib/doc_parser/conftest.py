#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""文档解析器测试 fixtures"""
import pytest
from pathlib import Path
from typing import List, Tuple


@pytest.fixture
def sample_docx_with_clauses(tmp_path: Path):
    """创建包含条款的 Word 文档 fixture"""
    pytest.importorskip("docx")
    from docx import Document

    def _create(docx_path: Path, clauses: List[Tuple[str, str, str]]) -> None:
        doc = Document()
        table = doc.add_table(rows=len(clauses) + 1, cols=2)
        table.rows[0].cells[0].text = "条款编号"
        table.rows[0].cells[1].text = "条款内容"
        for i, (number, title, text) in enumerate(clauses, 1):
            table.rows[i].cells[0].text = number
            table.rows[i].cells[1].text = f"{title}\n{text}"
        doc.save(str(docx_path))
    return _create


@pytest.fixture
def sample_docx_with_premium(tmp_path: Path):
    """创建包含费率表的 Word 文档 fixture"""
    pytest.importorskip("docx")
    from docx import Document

    def _create(docx_path: Path) -> None:
        doc = Document()
        table = doc.add_table(rows=4, cols=3)
        table.rows[0].cells[0].text = "年龄"
        table.rows[0].cells[1].text = "性别"
        table.rows[0].cells[2].text = "费率"
        for i, (age, gender, rate) in enumerate([
            ("18", "男", "100"),
            ("19", "男", "105"),
            ("20", "男", "110"),
        ], 1):
            table.rows[i].cells[0].text = age
            table.rows[i].cells[1].text = gender
            table.rows[i].cells[2].text = rate
        doc.save(str(docx_path))
    return _create


@pytest.fixture
def sample_docx_with_company_info(tmp_path: Path):
    """创建包含公司信息表格的 Word 文档 fixture"""
    pytest.importorskip("docx")
    from docx import Document

    def _create(docx_path: Path) -> None:
        doc = Document()
        table = doc.add_table(rows=3, cols=2)
        table.rows[0].cells[0].text = "公司名称"
        table.rows[0].cells[1].text = "XX保险公司"
        table.rows[1].cells[0].text = "地址"
        table.rows[1].cells[1].text = "北京市朝阳区..."
        table.rows[2].cells[0].text = "客服电话"
        table.rows[2].cells[1].text = "400-XXX-XXXX"
        doc.save(str(docx_path))
    return _create


@pytest.fixture
def sample_pdf_with_clauses(tmp_path: Path):
    """创建包含条款的 PDF 文档 fixture。集成测试建议使用真实文件。"""
    def _create(pdf_path: Path) -> None:
        pytest.importorskip("reportlab", reason="reportlab not installed")
        import reportlab.lib.pagesizes as pagesizes
        from reportlab.pdfgen import canvas
        c = canvas.Canvas(str(pdf_path), pagesize=pagesizes.A4)
        c.drawString(100, 700, "条款编号    条款内容")
        c.drawString(100, 680, "1           保险责任")
        c.save()
    return _create
